"""
Document Loader - Layer 1: Document Ingestion
Parses various document formats (PDF, DOCX, XLSX, TXT) into structured text + metadata.

M2 scope: PDF parser.
M3 scope: DOCX, XLSX, TXT parsers + DocumentLoader auto-detect class.
"""

import os
import logging

import fitz
import docx
import openpyxl

logger = logging.getLogger(__name__)


def parse_pdf(file_path):
    pages = []
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error(f"Failed to open PDF '{file_path}': {e}")
        return pages

    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            text = page.get_text().strip()
            pages.append({
                "page_num": page_index + 1,
                "text": text,
            })
    except Exception as e:
        logger.error(f"Failed while extracting text from '{file_path}': {e}")
    finally:
        doc.close()

    return pages


def parse_docx(file_path):
    sections = []
    try:
        document = docx.Document(file_path)
    except Exception as e:
        logger.error(f"Failed to open DOCX '{file_path}': {e}")
        return sections

    try:
        current_title = "Document"
        current_text_parts = []

        def flush_section():
            text = "\n".join(current_text_parts).strip()
            if text:
                sections.append({
                    "section_title": current_title,
                    "text": text,
                })

        for para in document.paragraphs:
            style_name = (para.style.name or "").lower()
            if style_name.startswith("heading") and para.text.strip():
                flush_section()
                current_title = para.text.strip()
                current_text_parts = []
            else:
                if para.text.strip():
                    current_text_parts.append(para.text.strip())

        flush_section()

        for t_index, table in enumerate(document.tables):
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            table_text = "\n".join(rows_text).strip()
            if table_text:
                sections.append({
                    "section_title": f"Table {t_index + 1}",
                    "text": table_text,
                })

    except Exception as e:
        logger.error(f"Failed while extracting text from '{file_path}': {e}")

    return sections


def parse_xlsx(file_path):
    results = []
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
    except Exception as e:
        logger.error(f"Failed to open XLSX '{file_path}': {e}")
        return results

    try:
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            header = [str(h) if h is not None else f"col{i}" for i, h in enumerate(rows[0])]
            row_lines = []

            for row in rows[1:]:
                if all(cell is None for cell in row):
                    continue
                pairs = []
                for col_name, value in zip(header, row):
                    if value is not None:
                        pairs.append(f"{col_name}: {value}")
                if pairs:
                    row_lines.append(" | ".join(pairs))

            sheet_text = "\n".join(row_lines).strip()
            if sheet_text:
                results.append({
                    "sheet_name": sheet.title,
                    "text": sheet_text,
                })
    except Exception as e:
        logger.error(f"Failed while extracting data from '{file_path}': {e}")

    return results


def parse_txt(file_path):
    encodings_to_try = ["utf-8", "latin-1"]

    for encoding in encodings_to_try:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
            return [{"text": text}] if text else []
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Failed to read TXT '{file_path}': {e}")
            return []

    logger.error(f"Could not decode '{file_path}' with any known encoding")
    return []


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md"}

    def load(self, file_path):
        file_name = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            msg = f"Unsupported file extension: '{ext}'"
            logger.warning(msg)
            return {
                "file_name": file_name,
                "file_type": ext.lstrip("."),
                "success": False,
                "chunks_raw": [],
                "error": msg,
            }

        if not os.path.exists(file_path):
            msg = f"File not found: '{file_path}'"
            logger.error(msg)
            return {
                "file_name": file_name,
                "file_type": ext.lstrip("."),
                "success": False,
                "chunks_raw": [],
                "error": msg,
            }

        try:
            if ext == ".pdf":
                chunks_raw = parse_pdf(file_path)
            elif ext == ".docx":
                chunks_raw = parse_docx(file_path)
            elif ext == ".xlsx":
                chunks_raw = parse_xlsx(file_path)
            elif ext in (".txt", ".md"):
                chunks_raw = parse_txt(file_path)
            else:
                chunks_raw = []

            return {
                "file_name": file_name,
                "file_type": ext.lstrip("."),
                "success": len(chunks_raw) > 0,
                "chunks_raw": chunks_raw,
                "error": None if len(chunks_raw) > 0 else "No extractable content found",
            }

        except Exception as e:
            msg = f"Unexpected error while loading '{file_path}': {e}"
            logger.error(msg)
            return {
                "file_name": file_name,
                "file_type": ext.lstrip("."),
                "success": False,
                "chunks_raw": [],
                "error": msg,
            }


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_loader.py <path_to_file>")
        sys.exit(1)

    loader = DocumentLoader()
    result = loader.load(sys.argv[1])

    print(f"File: {result['file_name']}")
    print(f"Type: {result['file_type']}")
    print(f"Success: {result['success']}")
    if result["error"]:
        print(f"Error: {result['error']}")
    print(f"Chunks extracted: {len(result['chunks_raw'])}")

    for c in result["chunks_raw"][:2]:
        preview_text = c.get("text", "")[:150].replace("\n", " ")
        label = c.get("page_num") or c.get("section_title") or c.get("sheet_name") or "-"
        print(f"\n--- {label} ---")
        print(f"{preview_text}...")
